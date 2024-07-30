import subprocess
import threading
import docx
import translator
import os


def translate_docx(threadsToUse: int, pdfName: str):
    outputDocPath = f'{pdfName}-output.docx'

    doc = docx.Document(outputDocPath)

    print(len(doc.paragraphs))
    threads = []

    paragraphs = len(doc.paragraphs)
    dividedParagraphs = paragraphs//threadsToUse

    paragraphsExcluded = 0
    for _, p in enumerate(doc.paragraphs):
        if p.text != "" and not translator.integer_checker.check_integer(p.text):
            paragraphsExcluded = paragraphsExcluded + 1

    for i in range(threadsToUse):
        if i == 0:
            start = 0
            end = dividedParagraphs
        else:
            start = dividedParagraphs*i
            end = start + dividedParagraphs + 1

        if i == threadsToUse-1:
            end = paragraphs

        thread = threading.Thread(target=translator.translate_paragraph, args=(doc, start, end, ))
        threads.append(thread)
        thread.start()

    for _, thread in enumerate(threads):
        thread.join()


    doc.save(f'./files/{pdfName}-outputTranslated.docx')


    subprocess.call(['soffice',
                    '--convert-to',
                    'pdf',
                    '--outdir',
                    './files',
                    f'./files/{pdfName}-outputTranslated.docx'])

    try:
        os.remove(outputDocPath)
    except:
        return
